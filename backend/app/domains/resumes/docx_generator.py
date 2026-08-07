import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def generate_docx_bytes(content: dict, template_name: str = "modern") -> bytes:
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    contact = content.get("contact", {})
    
    # Title / Name
    name_heading = doc.add_heading(contact.get("name", "Developer Resume"), level=0)
    name_heading.style.font.color.rgb = RGBColor(37, 99, 235)
    name_heading.style.font.size = Pt(20)

    # Contact line
    contact_bits = [b for b in [contact.get("email"), contact.get("phone"), contact.get("location"), contact.get("website")] if b]
    if contact_bits:
        cp = doc.add_paragraph(" • ".join(contact_bits))
        cp.style.font.size = Pt(9)
        cp.style.font.color.rgb = RGBColor(100, 116, 139)

    # Summary
    summary = content.get("summary")
    if summary:
        sh = doc.add_heading("SUMMARY", level=1)
        sh.style.font.size = Pt(12)
        doc.add_paragraph(summary)

    # Skills
    skills = content.get("skills", [])
    if skills:
        sk_h = doc.add_heading("SKILLS & TECH STACK", level=1)
        sk_h.style.font.size = Pt(12)
        doc.add_paragraph(", ".join(skills))

    # Experience
    experience = content.get("experience", [])
    if experience:
        exp_h = doc.add_heading("WORK EXPERIENCE", level=1)
        exp_h.style.font.size = Pt(12)
        for exp in experience:
            p = doc.add_paragraph()
            r1 = p.add_run(f"{exp.get('position')} ")
            r1.bold = True
            r2 = p.add_run(f"at {exp.get('company')} ")
            r2.italic = True
            p.add_run(f"({exp.get('start_date')} - {exp.get('end_date', 'Present')})")
            for h in exp.get("highlights", []):
                doc.add_paragraph(f"• {h}", style='List Bullet')

    # Projects
    projects = content.get("projects", [])
    if projects:
        pr_h = doc.add_heading("PROJECTS", level=1)
        pr_h.style.font.size = Pt(12)
        for proj in projects:
            p = doc.add_paragraph()
            r = p.add_run(proj.get("title"))
            r.bold = True
            if proj.get("description"):
                doc.add_paragraph(proj.get("description"))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
